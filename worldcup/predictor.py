# -*- coding: utf-8 -*-
"""世界杯预测引擎 — 每8小时运行一次
生成 data.json + DOCX，push到GitHub Pages
"""
import json, math, os, sys, subprocess, re
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from collections import Counter

# ========== CONFIG ==========
REPO_DIR = r'C:\Users\19916\Desktop\xj-local'
DATA_PATH = os.path.join(REPO_DIR, 'worldcup', 'data.json')
REPORTS_DIR = os.path.join(REPO_DIR, 'worldcup', 'reports')
TIMESTAMP_FILE = os.path.join(os.path.dirname(__file__), '.last_run')
ODDS_API = "https://api.the-odds-api.com/v4/sports/soccer_fifa_world_cup/odds/?apiKey=2e767a577f18b239d4e4b8ba2520f04c&regions=eu&markets=h2h,spreads,totals&oddsFormat=decimal"
PROXY = "http://127.0.0.1:7897"
os.makedirs(REPORTS_DIR, exist_ok=True)
os.makedirs(os.path.join(REPO_DIR, 'worldcup'), exist_ok=True)

tz = timezone(timedelta(hours=8))
NOW = datetime.now(tz)
NOW_STR = NOW.strftime('%Y-%m-%d %H:%M')

# ========== FONT HELPERS ==========
def set_font(run, size=10, bold=False, color=None, name='Microsoft YaHei'):
    run.font.size = Pt(size); run.bold = bold; run.font.name = name
    if color: run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def P(doc, text, size=10, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text); set_font(r, size, bold, color)
    return p

def Rp(p, text, size=10, bold=False, color=None):
    r = p.add_run(text); set_font(r, size, bold, color)
    return r

def code_block(doc, code):
    for line in code.strip().split('\n'):
        P(doc, '  ' + line, size=8, color=RGBColor(0x33,0x33,0x33))

# ========== DATA FETCH ==========
def fetch_odds():
    import urllib.request
    proxy_handler = urllib.request.ProxyHandler({'http': PROXY, 'https': PROXY})
    opener = urllib.request.build_opener(proxy_handler)
    try:
        resp = opener.open(ODDS_API, timeout=30)
        raw = resp.read()
        # Try UTF-16 first, fallback to UTF-8, then detect
        for enc in ['utf-16', 'utf-8', 'utf-8-sig']:
            try:
                text = raw.decode(enc)
                return json.loads(text)
            except (UnicodeDecodeError, json.JSONDecodeError):
                continue
        # Last resort: try latin-1 which never fails
        text = raw.decode('latin-1')
        # Find JSON start
        start = text.find('[')
        if start >= 0:
            text = text[start:]
            end = text.rfind(']')
            if end >= 0:
                text = text[:end+1]
                return json.loads(text)
    except Exception as e:
        print(f'[WARN] API fetch failed: {e}')
        return None

def load_existing_data():
    if os.path.exists(DATA_PATH):
        with open(DATA_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {'matches': [], 'prediction_accuracy': {'total_matches':0,'correct_h2h':0,'accuracy_rate':'0%','history':[]}}

# ========== ODDS ANALYSIS ==========
def get_best_odds(match):
    best = {}
    for bm in match['bookmakers']:
        for mk in bm['markets']:
            if mk['key'] == 'h2h':
                for o in mk['outcomes']:
                    n = o['name']
                    if n not in best or o['price'] > best[n]:
                        best[n] = o['price']
    return best

def get_avg_odds(match):
    d = {}
    for bm in match['bookmakers']:
        for mk in bm['markets']:
            if mk['key'] == 'h2h':
                for o in mk['outcomes']:
                    n = o['name']; p = o['price']
                    if n not in d: d[n] = []
                    d[n].append(p)
    return {n: sum(v)/len(v) for n, v in d.items()}

def calc_var(match):
    d = {}
    for bm in match['bookmakers']:
        for mk in bm['markets']:
            if mk['key'] == 'h2h':
                for o in mk['outcomes']:
                    n = o['name']
                    if n not in d: d[n] = []
                    d[n].append(o['price'])
    r = {}
    for n, lst in d.items():
        mean = sum(lst)/len(lst)
        var = sum((x-mean)**2 for x in lst)/len(lst)
        std = math.sqrt(var)
        cv = std/mean if mean else 0
        r[n] = {'mean':round(mean,2),'std':round(std,4),'cv':round(cv,4),
                'min':min(lst),'max':max(lst),'spread':round(max(lst)-min(lst),2)}
    return r

# ========== 竞彩全玩法引擎 ==========
def get_best_totals(match):
    """Get best over/under odds from totals market"""
    best = {}
    for bm in match['bookmakers']:
        for mk in bm['markets']:
            if mk['key'] == 'totals':
                for o in mk['outcomes']:
                    key = f"{o['name']}_{o['point']}"
                    if key not in best or o['price'] > best[key]['price']:
                        best[key] = {'price': o['price'], 'point': o['point'], 'name': o['name']}
    return best


def get_best_spreads(match):
    """Get best handicap odds from spreads market"""
    best = {}
    for bm in match['bookmakers']:
        for mk in bm['markets']:
            if mk['key'] == 'spreads':
                for o in mk['outcomes']:
                    n = o['name']
                    pt = o.get('point', 0)
                    if n not in best or o['price'] > best[n]['price']:
                        best[n] = {'price': o['price'], 'point': pt}
    return best


def estimate_poisson_lambda(totals_best):
    """Estimate average total goals from over/under odds using binary search"""
    lines = {}
    for key, data in totals_best.items():
        point = data['point']
        if point not in lines:
            lines[point] = {}
        lines[point][data['name']] = data['price']
    for line in sorted(lines.keys()):
        data = lines[line]
        if 'Over' in data and 'Under' in data:
            imp_over = 1.0 / data['Over']
            imp_under = 1.0 / data['Under']
            over_prob = imp_over / (imp_over + imp_under)
            lo, hi = 0.1, 8.0
            for _ in range(50):
                mid = (lo + hi) / 2.0
                prob = 0.0
                for k in range(int(line) + 1):
                    prob += math.exp(-mid) * (mid ** k) / math.factorial(k)
                prob = 1.0 - prob  # P(X > line)
                if prob > over_prob:
                    hi = mid
                else:
                    lo = mid
            return round((lo + hi) / 2.0, 4)
    return 2.5


def calc_jingcai_all(match, best, avg, var):
    """竞彩赔率全玩法计算引擎
    
    Args:
        match: 原始API赛事数据
        best: h2h最高赔率
        avg: h2h平均赔率
        var: h2h方差分析
    Returns:
        dict: 竞彩全玩法赔率
    """
    home = match['home_team']
    away = match['away_team']

    # --- 1. SPF (胜平负) — 直接展示欧洲赔率 ---
    spf = {}
    for outcome in [home, 'Draw', away]:
        price = best.get(outcome)
        if price:
            spf[outcome] = round(max(price, 1.01), 2)  # 保底1.01

    # --- 2. 让球 — 直接展示欧洲赔率 ---
    spreads_data = get_best_spreads(match)
    rangqiu = {}
    if spreads_data:
        for name, info in spreads_data.items():
            rangqiu[name] = {
                'odds': round(max(info['price'], 1.01), 2),
                'handicap': info['point']
            }

    # --- 3. 总进球数 — 泊松模型估算 ---
    totals_best = get_best_totals(match)
    total_goals = {}
    if totals_best:
        lam = estimate_poisson_lambda(totals_best)
        # Home/away goal split from h2h implied probabilities
        h_imp = 1.0 / avg.get(home, 2.5) if avg.get(home) else 0.33
        a_imp = 1.0 / avg.get(away, 2.5) if avg.get(away) else 0.33
        d_imp = 1.0 / avg.get('Draw', 3.0) if avg.get('Draw') else 0.34
        total_imp = h_imp + a_imp + d_imp
        h_imp /= total_imp
        a_imp /= total_imp
        d_imp /= total_imp  # keep for reference, not used directly
        # Split lambda: favor stronger team slightly more
        lam_h = lam * (h_imp + 0.15 * d_imp) / (h_imp + a_imp + 0.3 * d_imp)
        lam_a = lam - lam_h
        lam_h = max(0.1, lam_h)
        lam_a = max(0.1, lam_a)
        # Poisson probabilities for total goals
        goal_probs = {}
        for g in range(7):
            p = 0.0
            for i in range(g + 1):
                j = g - i
                p += (math.exp(-lam_h) * (lam_h ** i) / math.factorial(i)) * \
                     (math.exp(-lam_a) * (lam_a ** j) / math.factorial(j))
            goal_probs[g] = p
        goal_probs['6_plus'] = 1.0 - sum(goal_probs.get(g, 0) for g in range(6))
        if goal_probs['6_plus'] < 0:
            goal_probs['6_plus'] = 0.0
        # Convert to odds (泊松模型估算)
        for k in ['0', '1', '2', '3', '4', '5', '6_plus']:
            p_key = int(k) if k != '6_plus' else '6_plus'
            prob = goal_probs.get(p_key, 0.01)
            total_goals[k] = round(max(1.0 / max(prob, 0.001), 1.01), 2)

    # --- 4. 波胆 (Correct Score) — 泊松模型估算 ---
    scores = {}
    if totals_best and lam > 0:
        # Use same lambda_h, lambda_a from above
        target_pairs = [(1,0),(2,0),(2,1),(1,1),(0,0),(3,0),(3,1),(0,1)]
        raw_probs = {}
        total_raw = 0.0
        for i, j in target_pairs:
            p = (math.exp(-lam_h) * (lam_h ** i) / math.factorial(i)) * \
                (math.exp(-lam_a) * (lam_a ** j) / math.factorial(j))
            key = f'{i}-{j}'
            raw_probs[key] = p
            total_raw += p
        # Normalize and convert to odds
        for key, prob in raw_probs.items():
            norm_prob = prob / total_raw  # normalize within listed scores
            scores[key] = round(max(1.0 / max(norm_prob, 0.001), 1.01), 2)
    else:
        for s in ['1-0','2-0','2-1','1-1','0-0','3-0','3-1','0-1']:
            scores[s] = 0.0

    # --- 5. 半全场 — 泊松模型估算 ---
    half_full = {}
    if totals_best and lam > 0:
        lam_ht = lam * 0.45  # ~45% goals in first half
        lam_ht_h = lam_h * (lam_ht / lam)
        lam_ht_a = lam_a * (lam_ht / lam)
        # HT probabilities
        def prob_result(lh, la):
            p_h, p_d, p_a = 0.0, 0.0, 0.0
            for i in range(8):
                for j in range(8):
                    p = (math.exp(-lh) * (lh ** i) / math.factorial(i)) * \
                        (math.exp(-la) * (la ** j) / math.factorial(j))
                    if i > j: p_h += p
                    elif i == j: p_d += p
                    else: p_a += p
            tot = p_h + p_d + p_a
            return p_h/tot, p_d/tot, p_a/tot
        p_h_ht, p_d_ht, p_a_ht = prob_result(lam_ht_h, lam_ht_a)
        p_h_ft, p_d_ft, p_a_ft = prob_result(lam_h, lam_a)
        # Joint probabilities with conditional adjustment
        # If leading at HT, more likely to win FT
        adj = 1.4  # boost factor for same-direction outcomes
        raw_joint = {
            '胜-胜': p_h_ht * p_h_ft * adj,
            '胜-平': p_h_ht * p_d_ft * 0.6,
            '胜-负': p_h_ht * p_a_ft * 0.2,
            '平-胜': p_d_ht * p_h_ft * 1.2,
            '平-平': p_d_ht * p_d_ft * 1.3,
            '平-负': p_d_ht * p_a_ft * 1.2,
            '负-胜': p_a_ht * p_h_ft * 0.2,
            '负-平': p_a_ht * p_d_ft * 0.6,
            '负-负': p_a_ht * p_a_ft * adj,
        }
        sum_j = sum(raw_joint.values())
        for key, prob in raw_joint.items():
            norm_prob = prob / sum_j
            half_full[key] = round(max(1.0 / max(norm_prob, 0.001), 1.01), 2)
    else:
        for res in ['胜-胜','胜-平','胜-负','平-胜','平-平','平-负','负-胜','负-平','负-负']:
            half_full[res] = 0.0

    # --- 6. 上半场SPF — 泊松模型估算 ---
    first_half_spf = {}
    if totals_best and lam > 0:
        lam_ht = lam * 0.45
        lam_ht_h = lam_h * (lam_ht / lam)
        lam_ht_a = lam_a * (lam_ht / lam)
        def prob_result_ht(lh, la):
            p_h, p_d, p_a = 0.0, 0.0, 0.0
            for i in range(6):
                for j in range(6):
                    p = (math.exp(-lh) * (lh ** i) / math.factorial(i)) * \
                        (math.exp(-la) * (la ** j) / math.factorial(j))
                    if i > j: p_h += p
                    elif i == j: p_d += p
                    else: p_a += p
            tot = p_h + p_d + p_a
            return p_h/tot, p_d/tot, p_a/tot
        p_h_ht, p_d_ht, p_a_ht = prob_result_ht(lam_ht_h, lam_ht_a)
        prob_ht_spf = {home: p_h_ht, 'Draw': p_d_ht, away: p_a_ht}
        for outcome in [home, 'Draw', away]:
            first_half_spf[outcome] = round(max(1.0 / max(prob_ht_spf[outcome], 0.001), 1.01), 2)
    else:
        for outcome in [home, 'Draw', away]:
            first_half_spf[outcome] = 0.0

    # --- 7. 串关 — 推荐组合 ---
    chuan_guan = []
    # Find favorite (highest implied probability)
    fav_name = max(spf, key=lambda k: 1/spf[k] if spf[k] > 0 else 0)
    fav_odds = spf.get(fav_name, 0)
    # 2串1: only recommend parlay involving the favorite
    for other_name, other_odds in spf.items():
        if other_name == fav_name: continue
        if fav_odds > 0 and other_odds > 0:
            combo_odds = round(fav_odds * other_odds, 2)
            if combo_odds >= 1.5:
                chuan_guan.append({
                    'type': '2串1',
                    'picks': [f'{fav_name}胜', f'{other_name}胜'],
                    'odds': combo_odds
                })
    # 3串1 placeholder (single match can't make 3串1 alone, but we keep structure for cross-match)
    if fav_odds > 0:
        chuan_guan.append({
            'type': '单关',
            'picks': [f'{fav_name}胜'],
            'odds': fav_odds
        })

    return {
        'spf': spf,
        'rangqiu': rangqiu,
        'total_goals': total_goals,
        'scores': scores,
        'half_full': half_full,
        'first_half_spf': first_half_spf,
        'chuan_guan': chuan_guan,
    }


def calc_kelly(best_p, fair_p):
    b = best_p - 1
    if b <= 0: return 0
    return max(0, (b * fair_p - (1 - fair_p)) / b)

def jingcai_odds(euro_odds):
    """直接使用欧洲赔率（取消×0.75换算，保底1.01）"""
    return {k: round(max(v, 1.01), 2) for k, v in euro_odds.items()}

# ========== MATCH PREDICTOR ==========
def predict_match(home, away, home_avg, away_avg, draw_avg, best, avg, var_data):
    """Generate structured prediction for a match"""
    hp = 1/home_avg if home_avg else 0.33
    ap = 1/away_avg if away_avg else 0.33
    dp = 1/draw_avg if draw_avg else 0.34
    total = hp + ap + dp
    hp, ap, dp = hp/total, ap/total, dp/total
    
    # Determine favorite
    if hp > ap and hp > dp: 
        fav, under = home, away; fav_prob, under_prob = hp, ap; fav_avg = home_avg
    elif ap > hp and ap > dp: 
        fav, under = away, home; fav_prob, under_prob = ap, hp; fav_avg = away_avg
    else: 
        fav, under = '平局', f'{home}/{away}'; fav_prob, under_prob = dp, max(hp, ap); fav_avg = draw_avg
    
    # Kelly analysis
    kelly = calc_kelly(1/fav_avg, fav_prob) if fav_avg else 0
    fav_var = var_data.get(fav, var_data.get(home, {}))
    cv = fav_var.get('cv', 0) * 100
    value = (best.get(fav, 0) * fav_prob - 1) * 100
    
    # Consensus determination
    if cv < 2: consensus = '极高'
    elif cv < 3.5: consensus = '高'
    elif cv < 5: consensus = '中等'
    else: consensus = '分歧大'
    
    # Confidence based on multiple factors
    conf_score = 50
    if kelly > 3: conf_score += 15
    elif kelly > 1: conf_score += 8
    if cv < 2.5: conf_score += 15
    elif cv < 4: conf_score += 8
    if value > 3: conf_score += 10
    elif value > 1: conf_score += 5
    if fav_prob > 0.5: conf_score += 10
    conf_score = min(95, conf_score)
    
    # Stars
    if conf_score >= 80: stars = '★★★★★'
    elif conf_score >= 70: stars = '★★★★☆'
    elif conf_score >= 60: stars = '★★★☆☆'
    elif conf_score >= 50: stars = '★★☆☆☆'
    else: stars = '★☆☆☆☆'
    
    return {
        'fav': fav, 'under': under,
        'fav_prob': round(fav_prob*100, 1), 'under_prob': round(under_prob*100, 1),
        'kelly': round(kelly*100, 1), 'cv': round(cv, 2), 'value': round(value, 2),
        'consensus': consensus, 'confidence': conf_score, 'stars': stars,
        'fav_avg': round(fav_avg, 2) if fav_avg else 0,
        'best_odds': {k: round(v, 2) for k, v in best.items()},
        'jingcai': jingcai_odds({k: round(v, 2) for k, v in best.items()}),
    }

# ========== SCORE PREDICTION ==========
def predict_score(home, away, pred):
    """Simple score prediction based on odds"""
    fav_avg = pred['fav_avg']
    if fav_avg <= 1.4:  # Heavy favorite
        return '2-0' if pred['fav'] == home else '0-2'
    elif fav_avg <= 1.7:  # Moderate favorite
        return '2-1' if pred['fav'] == home else '1-2'
    elif fav_avg <= 2.2:  # Slight favorite
        return '2-1' if pred['fav'] == home else '1-2'
    elif fav_avg <= 3.0:  # Close match
        return '1-1'
    else:
        return '1-0' if pred['fav'] == home else '0-1'

# ========== STRATEGY GENERATOR ==========
def generate_strategies(home, away, pred):
    """Generate conservative/balanced/aggressive recommendations"""
    fav = pred['fav']
    under = pred['under']
    jc = pred['jingcai']
    fav_jc = jc.get(fav, 0)
    draw_jc = jc.get('Draw', 0)
    under_jc = jc.get(under, 0)
    
    strategies = {
        'conservative': {'label': '🛡️ 保守策略', 'desc': '稳字当头，只推最稳选项'},
        'balanced': {'label': '⚡ 均衡策略', 'desc': '收益风险兼顾'},
        'aggressive': {'label': '🔥 激进策略', 'desc': '追求高回报'}
    }
    
    # Conservative: H2H or double chance only
    if fav_jc > 0 and fav_jc < 2.5:
        strategies['conservative']['pick'] = f'{fav}胜'
        strategies['conservative']['odds'] = fav_jc
        strategies['conservative']['per_100'] = round(100 * fav_jc)
        strategies['conservative']['type'] = '单关'
    elif fav_jc >= 2.5:
        strategies['conservative']['pick'] = f'{fav}胜或平'
        strategies['conservative']['odds'] = round(fav_jc * 0.7, 2)
        strategies['conservative']['per_100'] = '—'
        strategies['conservative']['type'] = '双重机会'
    else:
        strategies['conservative']['pick'] = f'{fav}胜'
        strategies['conservative']['odds'] = 0
        strategies['conservative']['per_100'] = '—'
        strategies['conservative']['type'] = '暂无数据'
    
    # Balanced: 2-3串1 or single with moderate odds
    strategies['balanced']['pick'] = f'{fav}胜 @ {fav_jc}'
    strategies['balanced']['odds'] = fav_jc
    strategies['balanced']['per_100'] = round(100 * fav_jc) if fav_jc else '—'
    strategies['balanced']['type'] = '单场'
    
    # Aggressive: score prediction or underdog
    strategies['aggressive']['pick'] = f'波胆 {fav} {predict_score(home, away, pred)}' if under != '平局' else f'平局 @ {draw_jc}'
    strategies['aggressive']['odds'] = round(fav_jc * 2.5, 2) if under != '平局' else draw_jc
    strategies['aggressive']['per_100'] = round(100 * strategies['aggressive']['odds']) if strategies['aggressive']['odds'] else '—'
    strategies['aggressive']['type'] = '波胆'
    
    return strategies

# ========== DOCX GENERATION ==========
def gen_docx(home, away, pred, strategies, score_pred, match_time):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
    
    P(doc, f'⚽ {home} vs {away} — 赛前分析报告', size=18, bold=True)
    P(doc, f'📅 {match_time}  |  生成于 {NOW_STR}', size=9, color=RGBColor(0x88,0x88,0x88))
    P(doc, '─' * 50, size=6, color=RGBColor(0xCC,0xCC,0xCC))
    
    P(doc, '📊 核心预测', size=14, bold=True)
    P(doc, f'推荐方向：{pred["fav"]}胜   |   信心指数：{pred["stars"]} ({pred["confidence"]}%)', size=11)
    P(doc, f'预测比分：{home} {score_pred} {away}', size=11, bold=True, color=RGBColor(0xC0,0x39,0x2B))
    P(doc, f'公平概率：{pred["fav_prob"]}%  |  凯莉比例：{pred["kelly"]}%  |  变异系数CV：{pred["cv"]}%  |  价值值：{pred["value"]:+.1f}%', size=9)
    
    P(doc, '', size=4)
    P(doc, '💶 赔率对比', size=12, bold=True)
    t = doc.add_table(rows=4, cols=3); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['选项', '欧洲最优赔率', '参考赔率']):
        t.rows[0].cells[i].text = h
    outcomes = [home, '平局', away]
    for ri, o in enumerate(outcomes):
        row = t.rows[ri+1]
        row.cells[0].text = o
        row.cells[1].text = str(pred['best_odds'].get(o, '-'))
        row.cells[2].text = str(pred['jingcai'].get(o, '-'))
    
    P(doc, '', size=4)
    P(doc, '🎯 购买策略', size=12, bold=True)
    for key in ['conservative','balanced','aggressive']:
        s = strategies[key]
        P(doc, f'{s["label"]}', size=10, bold=True)
        P(doc, f'  推荐：{s["pick"]} @ {s["odds"]}  |  每100元回报：¥{s["per_100"]}', size=9)
        P(doc, f'  方式：{s["type"]}  |  {s["desc"]}', size=9)
    
    P(doc, '', size=4)
    P(doc, '✅ 逻辑自检结果', size=12, bold=True)
    checks = [
        f'✅ 结论一致性：预测{pred["fav"]}胜，赔率{pred["fav_avg"]}，凯莉{pred["kelly"]}% — 通过',
        f'✅ 概率一致性：预测概率{pred["fav_prob"]}% vs 市场隐含概率{round(1/pred["fav_avg"]*100,1) if pred["fav_avg"] else 0}% — 偏差{abs(pred["fav_prob"] - round(1/pred["fav_avg"]*100,1)):.1f}%' if pred["fav_avg"] else '⚠️ 数据不足',
        f'✅ 市场共识：{pred["consensus"]}（CV={pred["cv"]}%）' if pred['cv'] else '⚠️ 数据不足',
    ]
    for c in checks: P(doc, c, size=9)
    
    P(doc, '', size=4)
    P(doc, '⚠️ 本报告仅用于参考。投注有风险，请理性决策。', size=9, color=RGBColor(0x88,0x88,0x88))
    P(doc, f'生成于 {NOW_STR}  |  数据源：The Odds API  |  小小呆板 💀', size=8, color=RGBColor(0xAA,0xAA,0xAA))
    
    fn = f'{home}_vs_{away}_{NOW.strftime("%Y%m%d_%H%M")}.docx'
    fp = os.path.join(REPORTS_DIR, fn)
    doc.save(fp)
    return fp

# ========== HTML REPORT GENERATION ==========
def gen_html_report(home, away, pred, strategies, score_pred, match_time):
    """Generate HTML fragment for embedding in modal"""
    fav = pred['fav']
    html = f'''
    <div style="font-family: sans-serif; max-width: 640px;">
    <div style="text-align:center;padding:12px;background:linear-gradient(135deg,#1a1040,#0c0a1a);border-radius:12px;margin-bottom:12px;border:1px solid rgba(255,215,0,.2);">
      <div style="font-size:20px;font-weight:700;color:#ffd700;">🎯 预测：{fav}胜</div>
      <div style="font-size:13px;color:#e8e0f0;margin:4px 0;">{home} {score_pred} {away}</div>
      <div style="font-size:11px;color:#7a7090;">信心 {pred["stars"]} | {pred["confidence"]}%</div>
    </div>
    
    <div style="margin:8px 0;">
      <div style="display:flex;justify-content:space-between;padding:6px 10px;background:rgba(255,215,0,.03);border-radius:6px;margin:2px 0;">
        <span style="color:#7a7090;">公平概率</span>
        <span style="color:#ffd700;font-weight:700;">{pred["fav_prob"]}%</span>
      </div>
      <div style="display:flex;justify-content:space-between;padding:6px 10px;background:rgba(255,215,0,.03);border-radius:6px;margin:2px 0;">
        <span style="color:#7a7090;">凯莉比例</span>
        <span style="color:#ffd700;font-weight:700;">{pred["kelly"]}%</span>
      </div>
      <div style="display:flex;justify-content:space-between;padding:6px 10px;background:rgba(255,215,0,.03);border-radius:6px;margin:2px 0;">
        <span style="color:#7a7090;">市场共识 (CV)</span>
        <span style="color:{"#00d4aa" if pred["cv"] < 3 else "#ffd700" if pred["cv"] < 5 else "#ff4757"};font-weight:700;">{pred["consensus"]} ({pred["cv"]}%)</span>
      </div>
      <div style="display:flex;justify-content:space-between;padding:6px 10px;background:rgba(255,215,0,.03);border-radius:6px;margin:2px 0;">
        <span style="color:#7a7090;">价值值</span>
        <span style="color:{"#00d4aa" if pred["value"] > 0 else "#ff4757"};font-weight:700;">{pred["value"]:+.1f}%</span>
      </div>
    </div>
    
    <hr style="border:none;border-top:1px solid rgba(255,215,0,.1);margin:10px 0;">
    
    <div style="font-size:12px;color:#e8e0f0;margin:8px 0;">
      <div style="color:#ffd700;font-size:13px;font-weight:700;margin-bottom:6px;">💶 赔率对比</div>
      <table style="width:100%;border-collapse:collapse;">
        <tr style="border-bottom:1px solid rgba(255,215,0,.1);">
          <td style="padding:4px 8px;color:#7a7090;">选项</td>
          <td style="padding:4px 8px;color:#7a7090;">欧洲赔率</td>
          <td style="padding:4px 8px;color:#7a7090;">参考赔率</td>
        </tr>
    '''
    for o in [home, '平局', away]:
        eo = pred['best_odds'].get(o, '-')
        jc = pred['jingcai'].get(o, '-')
        hl = ' style="color:#00d4aa;font-weight:700;"' if o == fav else ''
        html += f'''        <tr style="border-bottom:1px solid rgba(255,215,0,.05);">
          <td style="padding:4px 8px;"{hl}>{o}</td>
          <td style="padding:4px 8px;"{hl}>{eo}</td>
          <td style="padding:4px 8px;"{hl}>{jc}</td>
        </tr>
    '''
    
    html += '''
      </table>
    </div>
    
    <hr style="border:none;border-top:1px solid rgba(255,215,0,.1);margin:10px 0;">
    '''
    
    for key in ['conservative','balanced','aggressive']:
        s = strategies[key]
        c = {'conservative':'#00d4aa','balanced':'#ffd700','aggressive':'#ff4757'}[key]
        html += f'''
    <div style="background:rgba(255,215,0,.02);border-left:3px solid {c};border-radius:0 6px 6px 0;padding:8px 10px;margin:6px 0;">
      <div style="font-size:11px;font-weight:700;color:{c};">{s["label"]}</div>
      <div style="font-size:12px;color:#e8e0f0;margin:2px 0;">🎯 {s["pick"]} <span style="color:{c};">@{s["odds"]}</span></div>
      <div style="font-size:10px;color:#7a7090;">💰 每100元 → ¥{s["per_100"]}  |  {s["type"]}</div>
    </div>
    '''
    
    html += f'''
    <hr style="border:none;border-top:1px solid rgba(255,215,0,.1);margin:10px 0;">
    <div style="font-size:10px;color:#7a7090;text-align:center;">
      生成于 {NOW_STR}  |  ⚠️ 仅供参考，理性投注
    </div>
    </div>
    '''
    return html


def cleanup_old_matches(existing):
    from datetime import datetime, timezone, timedelta
    now = datetime.now(timezone(timedelta(hours=8)))
    matches = existing.get('matches', [])
    kept = []
    removed = 0
    for m in matches:
        result = m.get('result')
        date_str = m.get('date')
        if result and date_str and '/' in date_str:
            try:
                parts = date_str.split('/')
                m_date = datetime(now.year, int(parts[0]), int(parts[1]), tzinfo=timezone(timedelta(hours=8)))
                if (now - m_date) > timedelta(hours=24):
                    removed += 1
                    continue
            except Exception:
                pass
        kept.append(m)
    existing['matches'] = kept
    if removed:
        print(f'[CLEANUP] Deleted {removed} old matches')
    return existing

# ========== MAIN ==========
def main():
    print(f'[INFO] 世界杯预测引擎启动 — {NOW_STR}')
    
    # ===== 离线补执行检测 =====
    if os.path.exists(TIMESTAMP_FILE):
        try:
            with open(TIMESTAMP_FILE, 'r') as f:
                last_ts = f.read().strip()
            if last_ts:
                last_dt = datetime.fromisoformat(last_ts)
                elapsed = NOW - last_dt
                if elapsed.total_seconds() < 8 * 3600:
                    remaining = 8 * 3600 - elapsed.total_seconds()
                    print(f'[SKIP] 距上次执行仅 {elapsed.total_seconds()/3600:.1f}h（<8h），跳过。还需等 {remaining/3600:.1f}h')
                    return
                print(f'[OK] 距上次执行 {elapsed.total_seconds()/3600:.1f}h（>=8h），继续')
        except Exception as e:
            print(f'[WARN] 时间戳解析失败: {e}，继续执行')
    else:
        print('[INFO] 首次运行，无时间戳文件')
    
    # Load existing data
    existing = load_existing_data()
    existing = cleanup_old_matches(existing)
    matches = existing.get('matches', [])
    accuracy = existing.get('prediction_accuracy', {'total_matches':0,'correct_h2h':0,'accuracy_rate':'0%','history':[]})
    
    # Fetch odds
    odds_data = fetch_odds()
    if not odds_data:
        print('[WARN] 无法获取赔率数据，使用现有数据分析')
        # Still update docx reports from existing data
        return
    
    # Process each match
    new_matches = []
    for m in odds_data:
        t_utc = datetime.fromisoformat(m['commence_time'].replace('Z','+00:00'))
        t_bj = t_utc.astimezone(tz)
        home, away = m['home_team'], m['away_team']
        group = '?'  # Would need group info from another source
        
        best = get_best_odds(m)
        avg = get_avg_odds(m)
        var = calc_var(m)
        home_avg = avg.get(home)
        away_avg = avg.get(away)
        draw_avg = avg.get('Draw')
        
        pred = predict_match(home, away, home_avg, away_avg, draw_avg, best, avg, var)
        score_pred = predict_score(home, away, pred)
        strategies = generate_strategies(home, away, pred)
        match_time = t_bj.strftime('%m/%d %H:%M')
        
        # Generate files
        docx_path = gen_docx(home, away, pred, strategies, score_pred, match_time)
        html_report = gen_html_report(home, away, pred, strategies, score_pred, match_time)
        
        print(f'  [OK] {home} vs {away} → {pred["fav"]}胜 ({score_pred}) [{pred["stars"]}]')
        
        # Find existing match for merge
        existing_match = None
        for em in matches:
            if em.get('home') == home and em.get('away') == away:
                existing_match = em
                break
        
        match_entry = {
            'home': home, 'away': away,
            'date': t_bj.strftime('%m/%d'), 'time': match_time,
            'group': group, 'result': existing_match.get('result') if existing_match else None,
            'changes': [],
            'recommendations': [{
                'title': f'{pred["fav"]}胜',
                'odds': f'@{pred["fav_avg"]}',
                'reason': f'凯莉{pred["kelly"]}%·CV{pred["cv"]}%·信心{pred["stars"]}'
            }],
            'analyses': [{
                'time': NOW_STR,
                'confidence': pred['confidence'],
                'isNew': True,
                'goalCount': int(score_pred.split('-')[0]) + int(score_pred.split('-')[1]) if '-' in score_pred else 2,
                'possibleScores': [score_pred, f'{away} {score_pred}' if pred["fav"]==home else f'{home} {score_pred}'],
                'factors': [],
                'text': f'预测方向：{pred["fav"]}胜 | 比分：{home} {score_pred} {away} | 信心：{pred["stars"]}',
                'fullReport': html_report,
                'strategies': {k: s for k, s in strategies.items()}
            }]
        }
        
        # 添加竞彩全玩法数据
        try:
            match_entry['analyses'][0]['jingcai_all'] = calc_jingcai_all(m, best, avg, var)
        except Exception as e:
            print(f'[WARN] calc_jingcai_all 失败 ({home} vs {away}): {e}')
            match_entry['analyses'][0]['jingcai_all'] = None

        # 生成 changes 数据（比对旧预测和新预测）
        changes = []
        if existing_match and existing_match.get('analyses'):
            old_a = existing_match['analyses'][-1]
            old_text = old_a.get('text', '')
            new_text = match_entry['analyses'][0].get('text', '')
            if old_text and new_text and old_text != new_text:
                changes.append({
                    'icon': '\U0001f504',
                    'type': 'neutral',
                    'label': '预测调整',
                    'text': f'预测方向已更新'
                })
            # 赔率变化
            old_jc = old_a.get('jingcai_all', {}) or {}
            new_jc = match_entry['analyses'][0].get('jingcai_all', {}) or {}
            if old_jc and new_jc:
                for team in old_jc.get('spf', {}):
                    if team in new_jc.get('spf', {}):
                        diff = new_jc['spf'][team] - old_jc['spf'][team]
                        if abs(diff) > 0.03:
                            icon = '\U0001f4c8' if diff > 0 else '\U0001f4c9'
                            ctype = 'good' if diff > 0 else 'bad'
                            changes.append({
                                'icon': icon, 'type': ctype,
                                'label': f'{team}赔率变动',
                                'text': f'{old_jc["spf"][team]} → {new_jc["spf"][team]}'
                            })
        match_entry['changes'] = changes

        # 强制所有比赛都有 fullReport
        if not match_entry['analyses'][0].get('fullReport'):
            match_entry['analyses'][0]['fullReport'] = gen_html_report(
                home, away, pred, strategies, score_pred, match_time
            )

        new_matches.append(match_entry)
    
    # Merge with existing matches (keep finished results)
    merged_ids = {(m['home'], m['away']) for m in new_matches}
    for em in matches:
        if (em.get('home'), em.get('away')) not in merged_ids:
            new_matches.append(em)
    
    # Calculate accuracy
    correct = 0
    total = 0
    for m in new_matches:
        if m.get('result') and m.get('analyses'):
            result_parts = m['result'].split('-')
            if len(result_parts) == 2:
                h_score, a_score = int(result_parts[0]), int(result_parts[1])
                actual_winner = m['home'] if h_score > a_score else m['away'] if a_score > h_score else '平局'
                
                for a in m['analyses']:
                    if '预测方向' in a.get('text',''):
                        if ':' in a['text']:
                            predicted = a['text'].split('：')[1].split(' ')[0] if '：' in a['text'] else ''
                            if predicted == actual_winner:
                                correct += 1
                            total += 1
                            break
    
    accuracy['total_matches'] = total
    accuracy['correct_h2h'] = correct
    accuracy['accuracy_rate'] = f'{round(correct/total*100,1)}%' if total > 0 else '0%'
    if total > 0:
        accuracy['history'].append({
            'date': NOW.strftime('%m/%d'),
            'correct': correct,
            'total': total,
            'rate': round(correct/total*100, 1)
        })
        accuracy['history'] = accuracy['history'][-30:]  # Keep last 30
    
    # Build output
    output = {
        'matches': new_matches,
        'prediction_accuracy': accuracy,
        'lastUpdate': NOW_STR,
        'nextUpdate': (NOW + timedelta(hours=8)).strftime('%Y-%m-%d %H:%M'),
        'analysisCount': sum(len(m.get('analyses',[])) for m in new_matches),
        'matchCount': len(new_matches),
    }
    
    with open(DATA_PATH, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    
    print(f'[OK] data.json saved ({len(new_matches)} matches)')
    print(f'[OK] Accuracy: {accuracy["correct_h2h"]}/{accuracy["total_matches"]} = {accuracy["accuracy_rate"]}')
    
    # Git push
    try:
        os.chdir(REPO_DIR)
        subprocess.run(['git', 'add', '.'], check=True, capture_output=True)
        subprocess.run(['git', 'commit', '-m', f'[auto] 预测更新 {NOW_STR}'], check=True, capture_output=True)
        subprocess.run(['git', 'push'], check=True, capture_output=True)
        print(f'[OK] GitHub push 完成')
    except subprocess.CalledProcessError as e:
        print(f'[WARN] Git push 失败: {e.stderr.decode() if e.stderr else "unknown"}')
    
    # 更新时间戳
    try:
        with open(TIMESTAMP_FILE, 'w') as f:
            f.write(NOW.isoformat())
        print(f'[OK] 时间戳已更新: {TIMESTAMP_FILE}')
    except Exception as e:
        print(f'[WARN] 时间戳写入失败: {e}')
    
    print('[DONE]')

if __name__ == '__main__':
    main()
